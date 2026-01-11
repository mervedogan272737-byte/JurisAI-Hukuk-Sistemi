import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from docx import Document
from io import BytesIO
from PyPDF2 import PdfReader

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="JurisAI | Hukuki Karar Destek", page_icon="⚖️", layout="wide")

# --- GÖRSEL TASARIM (CSS) ---
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .report-card { padding: 25px; border-radius: 15px; background-color: white; border-left: 5px solid #1f77b4; box-shadow: 0 4px 6px rgba(0,0,0,0.1); color: #333; }
    .stButton>button { width: 100%; border-radius: 10px; background-color: #1f77b4; color: white; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

# --- WORD DOSYASI FONKSİYONU ---
def word_hazirla(baslik, icerik):
    doc = Document()
    doc.add_heading(baslik, 0)
    doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d/%m/%Y')}\n")
    doc.add_paragraph(icerik)
    doc.add_paragraph("\n\nSaygılarımla,\nAv. Merve Kılıç")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- GİRİŞ KONTROLÜ ---
if "login" not in st.session_state:
    st.session_state.login = False

if not st.session_state.login:
    st.title("⚖️ JurisAI Avukat Girişi")
    sifre = st.text_input("Sistem Şifresini Giriniz", type="password")
    if st.button("Giriş Yap"):
        if sifre == "avk2026": # Şifren bu!
            st.session_state.login = True
            st.rerun()
        else:
            st.error("Hatalı Şifre!")
else:
    # --- ANA PANEL ---
    with st.sidebar:
        st.title("JurisAI v2.0")
        st.info("Lisans: Merve Kılıç")
        secim = st.radio("Menü", ["📊 Analiz Merkezi", "📝 Dilekçe Yazıcı"])
        if st.button("Çıkış"):
            st.session_state.login = False
            st.rerun()

    if secim == "📊 Analiz Merkezi":
        st.title("📊 Dava Strateji Analizi")
        guc = st.slider("Kanıt Gücü (%)", 0, 100, 70)
        df = pd.DataFrame(dict(r=[guc, 80, 60, 90, 75], theta=['Kanıt', 'Yargıtay', 'Süreç', 'Kazanma', 'Verim']))
        fig = px.line_polar(df, r='r', theta='theta', line_close=True)
        st.plotly_chart(fig)

    else:
        st.title("📝 Akıllı Dilekçe Yazıcı")
        
        # PDF Yükleme Eksikliği Giderildi
        pdf_dosya = st.file_uploader("Analiz edilecek dosya (PDF)", type=["pdf"])
        pdf_metin = ""
        if pdf_dosya:
            reader = PdfReader(pdf_dosya)
            pdf_metin = "\n".join([p.extract_text() for p in reader.pages])
            st.success("PDF Verisi Alındı!")

        t_ad = st.text_input("Müvekkil Ad Soyad")
        d_tipi = st.selectbox("Tür", ["İhtarname", "Dava Dilekçesi", "Cevap Dilekçesi"])

        if st.button("Taslağı Oluştur"):
            if t_ad:
                metin = f"{t_ad} vekili olarak; ekli belgeler ve mevzuat uyarınca haklarımızın teminini talep ederiz."
                if pdf_metin:
                    metin += f"\n\nDosya Analiz Notu: {pdf_metin[:300]}..."

                # Ekranda Şık Gösterim
                st.markdown(f"<div class='report-card'><h3>{d_tipi.upper()}</h3><p>{metin}</p></div>", unsafe_allow_html=True)
                
                # Word İndirme Eksikliği Giderildi
                w_buf = word_hazirla(d_tipi, metin)
                st.download_button("📄 Word Olarak İndir", w_buf, f"{t_ad}.docx")
                
                # Uyarı artık sadece en altta küçük not olarak var
                st.caption("⚠️ Bu bir yapay zeka taslağıdır, avukat kontrolü gereklidir.")
